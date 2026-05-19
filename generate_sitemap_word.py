#!/usr/bin/env python3
"""
Generate NanoBio Studio Sitemap as Word Document (Landscape)
Updated with latest improvements (March 19, 2026)
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

# Create Document
doc = Document()

# Set to landscape orientation
section = doc.sections[0]
section.page_height = Inches(8.5)
section.page_width = Inches(11)
section.orientation = WD_ORIENT.LANDSCAPE
section.top_margin = Inches(0.7)
section.bottom_margin = Inches(0.7)
section.left_margin = Inches(0.6)
section.right_margin = Inches(0.6)

# Color definitions (RGB)
DARK_BLUE = RGBColor(26, 35, 126)      # #1a237e
PURPLE_BLUE = RGBColor(102, 126, 234)  # #667eea
LIGHT_GRAY = RGBColor(240, 240, 240)   # #f0f0f0
VERY_LIGHT_GRAY = RGBColor(249, 249, 249)  # #f9f9f9

def add_title(doc, title):
    """Add a title paragraph"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = DARK_BLUE
    return p

def add_subtitle(doc, subtitle):
    """Add a subtitle paragraph"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = PURPLE_BLUE

def add_section_heading(doc, heading):
    """Add a section heading"""
    p = doc.add_paragraph()
    run = p.add_run(heading)
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = DARK_BLUE
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(10)

def add_subsection_heading(doc, heading):
    """Add a subsection heading"""
    p = doc.add_paragraph()
    run = p.add_run(heading)
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = PURPLE_BLUE
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

def add_bullet_point(doc, text, indent=0):
    """Add a bullet point"""
    if indent == 0:
        p = doc.add_paragraph(text, style='List Bullet')
    else:
        p = doc.add_paragraph(text, style='List Bullet 2')
    
    for run in p.runs:
        run.font.size = Pt(9)
    p.paragraph_format.space_after = Pt(4)

def shade_table_header(table):
    """Shade the header row of a table"""
    for cell in table.rows[0].cells:
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), '1a237e')
        cell._element.get_or_add_tcPr().append(shading_elm)
        
        # Set text color to white
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True

def shade_alternating_rows(table, color1='f0f0f0', color2='ffffff'):
    """Shade alternating rows"""
    for i, row in enumerate(table.rows[1:], 1):
        color = color1 if i % 2 == 1 else color2
        for cell in row.cells:
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), color)
            cell._element.get_or_add_tcPr().append(shading_elm)

def set_table_borders(table):
    """Set table borders"""
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'cccccc')
        tblBorders.append(border)
    
    tblPr.append(tblBorders)

# ============================================================================
# TITLE PAGE
# ============================================================================

add_title(doc, "🏗️ NanoBio Studio")
add_subtitle(doc, "Application Site Map & Architecture")
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
run.font.size = Pt(10)
run.font.italic = True

doc.add_paragraph()

# Info table
info_table = doc.add_table(rows=4, cols=2)
info_table.style = 'Light Grid Accent 1'
shade_table_header(info_table)
set_table_borders(info_table)

info_data = [
    ("App Type:", "Streamlit Multi-Page Application"),
    ("Status:", "Production Ready"),
    ("Last Update:", "March 19, 2026"),
    ("Database:", "SQLite (trial_registry.db, nanobio_studio.db)"),
]

for i, (label, value) in enumerate(info_data):
    row = info_table.rows[i]
    row.cells[0].text = label
    row.cells[1].text = value
    
    for run in row.cells[0].paragraphs[0].runs:
        run.font.bold = True
        run.font.size = Pt(9)
    
    for run in row.cells[1].paragraphs[0].runs:
        run.font.size = Pt(9)

shade_alternating_rows(info_table)

doc.add_page_break()

# ============================================================================
# TABLE OF CONTENTS
# ============================================================================

add_section_heading(doc, "📑 Table of Contents")

toc_items = [
    "1. Authentication & Session Management",
    "2. Main Workflow Pages",
    "3. AI Co-Designer Services",
    "4. Education & Training",
    "5. Administration & Data Management",
    "6. User Roles & Access Control",
    "7. Database Architecture",
    "8. Recent Improvements (March 2026)",
    "9. Key Features & Capabilities",
]

for item in toc_items:
    add_bullet_point(doc, item)

doc.add_page_break()

# ============================================================================
# 1. AUTHENTICATION & SESSION MANAGEMENT
# ============================================================================

add_section_heading(doc, "1️⃣ Authentication & Session Management")

auth_items = [
    "Entry Point: Login.py",
    "Technology: Streamlit Auth with Persistent Session Storage",
    "Status: ✅ Fully implemented with 15-minute idle timeout",
]

for item in auth_items:
    p = doc.add_paragraph(item)
    for run in p.runs:
        run.font.size = Pt(9)

add_subsection_heading(doc, "Session Features:")

session_features = [
    "✅ Persistent login (survives page refresh)",
    "✅ 15-minute idle timeout with auto-logout",
    "✅ Session tokens stored in sessions.json",
    "✅ Session data: username, user_id, email, roles, timestamps",
    "✅ Token passing via URL query parameters (?session_token=...)",
    "✅ Auto-restoration of session on page load",
]

for feat in session_features:
    add_bullet_point(doc, feat)

doc.add_paragraph()

# ============================================================================
# 2. MAIN WORKFLOW PAGES
# ============================================================================

add_section_heading(doc, "2️⃣ Main Workflow Pages")

workflow_table = doc.add_table(rows=11, cols=4)
workflow_table.style = 'Light Grid Accent 1'
set_table_borders(workflow_table)

workflow_data = [
    ["Page", "Route", "Purpose", "Features"],
    ["🔐 Login", "Login.py", "Authentication Entry Point", "Persistent session, RBAC check"],
    ["0️⃣ Disease Selection", "pages/0_Disease_Selection.py", "Workflow Step 1", "Disease & pathology selection"],
    ["1️⃣ Design Parameters", "pages/1_Design_Parameters.py", "Workflow Step 2", "NP design parameters"],
    ["2️⃣ Run Simulation", "pages/2_Run_Simulation.py", "Workflow Step 3 (FINAL)", "Runs simulation, saves to DB"],
    ["6️⃣ Trial History", "pages/6_Trial_History.py", "Advanced Feature", "Loads from trial_registry.db"],
    ["7️⃣ AI Co-Designer", "pages/7_AI_Co_Designer.py", "AI Services", "AI-powered suggestions"],
    ["8️⃣ Protocol Generator", "pages/8_Protocol_Generator.py", "Lab Protocols", "Generate synthesis protocols"],
    ["9️⃣ AI Architecture", "pages/9_AI_Architecture.py", "System Design", "Architecture documentation"],
    ["🔟 ML Training", "pages/10_ML_Training.py", "Machine Learning", "Model training interface"],
]

for i, row_data in enumerate(workflow_data):
    row = workflow_table.rows[i]
    for j, cell_text in enumerate(row_data):
        row.cells[j].text = cell_text
        for run in row.cells[j].paragraphs[0].runs:
            run.font.size = Pt(8)

shade_table_header(workflow_table)
shade_alternating_rows(workflow_table)

doc.add_paragraph()

add_subsection_heading(doc, "All Pages Have:")
page_features = [
    "✅ Redirect button when not logged in ('Go to Login' with st.switch_page())",
    "✅ Timeout detection with auto-logout after 30 minutes",
    "✅ Clear messaging about session status",
]

for feat in page_features:
    add_bullet_point(doc, feat)

doc.add_page_break()

# ============================================================================
# 3. AI CO-DESIGNER SERVICES
# ============================================================================

add_section_heading(doc, "3️⃣ AI Co-Designer Services")

ai_table = doc.add_table(rows=3, cols=3)
ai_table.style = 'Light Grid Accent 1'
set_table_borders(ai_table)

ai_data = [
    ["Service", "Route", "Purpose"],
    ["AI Co-Designer", "pages/7_AI_Co_Designer.py", "AI-powered design suggestions & optimization"],
    ["Documentation", "pages/About_AI_Co_Designer.py", "AI system documentation & capabilities"],
]

for i, row_data in enumerate(ai_data):
    row = ai_table.rows[i]
    for j, cell_text in enumerate(row_data):
        row.cells[j].text = cell_text
        for run in row.cells[j].paragraphs[0].runs:
            run.font.size = Pt(9)

shade_table_header(ai_table)
shade_alternating_rows(ai_table)

doc.add_page_break()

# ============================================================================
# 4. EDUCATION & TRAINING
# ============================================================================

add_section_heading(doc, "4️⃣ Education & Training")

training_table = doc.add_table(rows=4, cols=3)
training_table.style = 'Light Grid Accent 1'
set_table_borders(training_table)

training_data = [
    ["Component", "Route", "Purpose"],
    ["Tutorial", "pages/10_Tutorial.py", "Learning resources & guides"],
    ["ML Training (5 pages)", "pages/13_ML Training/", "ML architecture & model training"],
    ["Architecture", "pages/11_AI_Architecture.py", "System architecture documentation"],
]

for i, row_data in enumerate(training_data):
    row = training_table.rows[i]
    for j, cell_text in enumerate(row_data):
        row.cells[j].text = cell_text
        for run in row.cells[j].paragraphs[0].runs:
            run.font.size = Pt(9)

shade_table_header(training_table)
shade_alternating_rows(training_table)

doc.add_paragraph()

add_subsection_heading(doc, "ML Training Pages:")
ml_pages = [
    "1_AI_Architecture.py - AI system design",
    "2_Model_Training_Process.py - Model training workflow",
    "3_Feature_Engineering.py - Feature engineering techniques",
    "4_Validation_Testing.py - Model validation & testing",
    "5_Dataset_Statistics.py - Data analysis & statistics",
]

for page in ml_pages:
    add_bullet_point(doc, page)

doc.add_page_break()

# ============================================================================
# 5. ADMINISTRATION & DATA MANAGEMENT
# ============================================================================

add_section_heading(doc, "5️⃣ Administration & Data Management")

admin_items = [
    "Admin Dashboard: audit_dashboard.py",
    "User Management: users.json (admin, designer, viewer roles)",
    "Database Management: Multiple utility scripts for data maintenance",
]

for item in admin_items:
    p = doc.add_paragraph(item)
    for run in p.runs:
        run.font.size = Pt(9)

add_subsection_heading(doc, "Admin Tools:")
admin_tools = [
    "✅ Audit dashboard for system monitoring",
    "✅ User role management (Admin, Designer, Viewer)",
    "✅ Dataset management & import",
    "✅ Model storage & versioning",
    "✅ Data source integration (ToxCast API, external sources)",
]

for tool in admin_tools:
    add_bullet_point(doc, tool)

doc.add_page_break()

# ============================================================================
# 6. USER ROLES & ACCESS CONTROL
# ============================================================================

add_section_heading(doc, "6️⃣ User Roles & Access Control")

roles_table = doc.add_table(rows=4, cols=3)
roles_table.style = 'Light Grid Accent 1'
set_table_borders(roles_table)

roles_data = [
    ["Role", "Permissions", "Restrictions"],
    ["👤 Admin", "Full access to all pages, user management, data sources, audit dashboard", "None"],
    ["👨‍💼 Designer", "Workflow pages (0-17), AI Co-Designer, Trial History, Tutorial", "No user mgmt, no audit"],
    ["👁️ Viewer", "Read-only access to Trial History, Analytics, Tutorial", "No design/edit perms"],
]

for i, row_data in enumerate(roles_data):
    row = roles_table.rows[i]
    for j, cell_text in enumerate(row_data):
        row.cells[j].text = cell_text
        for run in row.cells[j].paragraphs[0].runs:
            run.font.size = Pt(8)

shade_table_header(roles_table)
shade_alternating_rows(roles_table)

doc.add_page_break()

# ============================================================================
# 7. DATABASE ARCHITECTURE
# ============================================================================

add_section_heading(doc, "7️⃣ Database Architecture")

add_subsection_heading(doc, "Primary Databases:")

db_table = doc.add_table(rows=4, cols=3)
db_table.style = 'Light Grid Accent 1'
set_table_borders(db_table)

db_data = [
    ["Database", "Tables", "Purpose"],
    ["trial_registry.db", "trials, trial_sequences", "Stores all simulation trials & results"],
    ["nanobio_studio.db", "designs, parameters, results", "Design configurations & outputs"],
    ["users.db", "users, roles", "User authentication & RBAC"],
]

for i, row_data in enumerate(db_data):
    row = db_table.rows[i]
    for j, cell_text in enumerate(row_data):
        row.cells[j].text = cell_text
        for run in row.cells[j].paragraphs[0].runs:
            run.font.size = Pt(9)

shade_table_header(db_table)
shade_alternating_rows(db_table)

doc.add_paragraph()

add_subsection_heading(doc, "Session Storage:")
storage_items = [
    "• sessions.json - Persistent user sessions with tokens & timestamps",
    "• users.json - User credentials & metadata",
]

for item in storage_items:
    p = doc.add_paragraph(item)
    for run in p.runs:
        run.font.size = Pt(9)

doc.add_paragraph()

add_subsection_heading(doc, "Key Tables (trial_registry.db):")
trial_schema = [
    "trials: trial_id, disease_subtype, disease_name, drug_name, np_size_nm,",
    "        np_charge_mv, np_peg_percent, np_zeta_potential, np_pdi,",
    "        treatment_dose_mgkg, treatment_route, treatment_frequency,",
    "        treatment_duration_days, trial_outcomes, creation_timestamp, status, notes",
    "",
    "trial_sequences: date, disease_code, next_sequence",
]

for line in trial_schema:
    p = doc.add_paragraph(line)
    for run in p.runs:
        run.font.size = Pt(8)
        run.font.name = 'Courier New'

doc.add_page_break()

# ============================================================================
# 8. RECENT IMPROVEMENTS
# ============================================================================

add_section_heading(doc, "8️⃣ Recent Improvements (March 2026)")

add_subsection_heading(doc, "Session Persistence System (March 18-19):")
improvements = [
    "✅ Implemented persistent login surviving page refresh",
    "✅ Added 15-minute idle timeout with auto-logout",
    "✅ Session tokens stored in sessions.json with timestamps",
    "✅ Session restoration on page load via URL query parameters",
    "✅ Clear messaging on timeout with redirect button",
]

for imp in improvements:
    add_bullet_point(doc, imp)

add_subsection_heading(doc, "Redirect Button Implementation:")
redirect_imp = [
    "✅ Added 'Go to Login' button to all protected pages",
    "✅ Button uses st.switch_page() for explicit navigation",
    "✅ Applied to 11+ pages across entire application",
    "✅ Consistent UX: message + centered button + st.stop()",
    "✅ Buttons clear query params before redirect",
]

for imp in redirect_imp:
    add_bullet_point(doc, imp)

add_subsection_heading(doc, "Trial History Fix:")
trial_imp = [
    "✅ Trials now saved to persistent trial_registry.db",
    "✅ Run Simulation calls create_trial_entry() for each new trial",
    "✅ Trial History reads from database via get_all_trials()",
    "✅ Trial count no longer stuck at 31 after refresh",
    "✅ New trials persist across sessions & user logins",
]

for imp in trial_imp:
    add_bullet_point(doc, imp)

doc.add_page_break()

# ============================================================================
# 9. KEY FEATURES & CAPABILITIES
# ============================================================================

add_section_heading(doc, "9️⃣ Key Features & Capabilities")

add_subsection_heading(doc, "Core Features:")
features = [
    "🧪 Nanoparticle Design Optimization - Multi-parameter design tool",
    "📊 Simulation Engine - Kinetics, biodistribution, safety predictions",
    "🤖 AI Co-Designer - AI-powered design suggestions",
    "📈 Trial History - Track & compare all designs",
    "📋 Data Analytics - Performance trends & metrics",
    "🔐 Authentication - Multi-user support with RBAC",
    "⏱️ Session Management - Persistent login with idle timeout",
    "💾 Data Persistence - SQLite databases for trials & configurations",
]

for feat in features:
    add_bullet_point(doc, feat)

add_subsection_heading(doc, "Project Structure:")
structure = [
    "/pages/ → Main application pages (0-17 + utilities)",
    "/pages/13_ML Training/ → 5-page ML training module",
    "/ai_engine/ → AI optimization & analysis",
    "/core/ → Scoring, analysis, algorithms",
    "/utils/ → PDF generation, helpers",
    "/modules/ → Trial registry, clinical data",
    "/tabs/ → Internal tab components",
    "/components/ → UI components & utilities",
]

for item in structure:
    add_bullet_point(doc, item)

doc.add_page_break()

# ============================================================================
# FOOTER
# ============================================================================

footer_table = doc.add_table(rows=5, cols=3)
footer_table.style = 'Light Grid Accent 1'
set_table_borders(footer_table)

footer_data = [
    ["", "", ""],
    ["Repository", "NanoBio Studio (Private)", "GitHub"],
    ["Deployment", "Streamlit Cloud", "Automatic via git push"],
    ["Version", "1.0.0", "Production"],
    ["Last Deploy", f"{datetime.now().strftime('%B %d, %Y')}", "Successful"],
]

for i, row_data in enumerate(footer_data):
    row = footer_table.rows[i]
    for j, cell_text in enumerate(row_data):
        row.cells[j].text = cell_text
        for run in row.cells[j].paragraphs[0].runs:
            run.font.size = Pt(9)

shade_table_header(footer_table)
shade_alternating_rows(footer_table)

doc.add_paragraph()

footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer_p.add_run("📧 For questions or updates, contact the development team")
run.font.size = Pt(8)
run.font.italic = True
run.font.color.rgb = RGBColor(128, 128, 128)

# ============================================================================
# SAVE DOCUMENT
# ============================================================================

word_path = "SITEMAP_LATEST.docx"
doc.save(word_path)
print(f"✅ Word document generated successfully: {word_path}")
print(f"   Format: Landscape")
print(f"   Orientation: {section.orientation}")
